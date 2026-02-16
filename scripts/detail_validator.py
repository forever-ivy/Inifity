#!/usr/bin/env python3
"""Structure-based format validation for translated documents.

Compares original vs translated documents at XML/property level (not visual).
Supports DOCX (python-docx) and XLSX (openpyxl).

Generates:
1. JSON validation report with issues and scores
2. Markdown report for LLM consumption with fix hints
"""

from __future__ import annotations

import dataclasses
import json
import os
from dataclasses import dataclass, field
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Any, Callable

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Border, Alignment
    OPENPYXL_AVAILABLE = True
except Exception:
    OPENPYXL_AVAILABLE = False


class Severity(str, Enum):
    """Issue severity levels."""
    CRITICAL = "error"
    WARNING = "warning"
    INFO = "info"


class Category(str, Enum):
    """Issue categories for DOCX."""
    # DOCX categories
    FONT = "font"
    STYLE = "style"
    TABLE = "table"
    TABLE_BORDERS = "table_borders"
    PARAGRAPH = "paragraph"
    HEADER_FOOTER = "header_footer"
    LIST = "list"
    NUMBERING = "numbering"
    # XLSX categories
    CELL_FONT = "cell_font"
    CELL_BORDER = "cell_border"
    CELL_FILL = "cell_fill"
    CELL_ALIGNMENT = "cell_alignment"
    MERGED_REGIONS = "merged_regions"
    DIMENSIONS = "dimensions"
    NUMBER_FORMAT = "number_format"


@dataclass(frozen=True)
class ValidationIssue:
    """A single validation finding."""
    category: str
    severity: str
    location: str
    element_type: str
    expected: str
    actual: str
    hint: str
    context: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "severity": self.severity,
            "category": self.category,
            "location": self.location,
            "element_type": self.element_type,
            "expected": self.expected,
            "actual": self.actual,
            "fix_hint": self.hint,
            **self.context,
        }


@dataclass
class ValidationResult:
    """Validation result for a single file."""
    file_name: str
    file_path: str
    format_type: str
    valid: bool
    total_checks: int = 0
    passed: int = 0
    warnings: int = 0
    failed: int = 0
    issues: list[ValidationIssue] = field(default_factory=list)
    format_fidelity_score: float = 1.0
    validation_timestamp: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

    def add_issue(self, issue: ValidationIssue) -> None:
        self.issues.append(issue)
        if issue.severity == Severity.CRITICAL:
            self.failed += 1
        elif issue.severity == Severity.WARNING:
            self.warnings += 1
        self.total_checks += 1

    def add_pass(self) -> None:
        self.passed += 1
        self.total_checks += 1

    def calculate_score(self, critical_weight: float = 0.15, warning_weight: float = 0.05) -> None:
        """Calculate format fidelity score based on issues."""
        deductions = self.failed * critical_weight + self.warnings * warning_weight
        self.format_fidelity_score = max(0.0, 1.0 - deductions)
        self.valid = self.format_fidelity_score >= 0.85 and self.failed == 0

    def to_dict(self) -> dict[str, Any]:
        return {
            "file_name": self.file_name,
            "file_path": self.file_path,
            "format_type": self.format_type,
            "valid": self.valid,
            "total_checks": self.total_checks,
            "passed": self.passed,
            "warnings": self.warnings,
            "failed": self.failed,
            "score": round(self.format_fidelity_score, 3),
            "issues": [i.to_dict() for i in self.issues],
        }


@dataclass
class ValidationConfig:
    """Configuration for detail validation."""
    # Enable/disable validators
    enable_docx: bool = True
    enable_xlsx: bool = True

    # DOCX settings
    docx_check_fonts: bool = True
    docx_check_styles: bool = True
    docx_check_tables: bool = True
    docx_check_headers_footers: bool = True
    docx_check_lists: bool = True
    docx_font_size_delta: float = 0.5  # pt tolerance

    # XLSX settings
    xlsx_check_fonts: bool = True
    xlsx_check_borders: bool = True
    xlsx_check_fills: bool = True
    xlsx_check_alignment: bool = True
    xlsx_check_merged: bool = True
    xlsx_check_dimensions: bool = True
    xlsx_check_number_formats: bool = True
    xlsx_dimension_delta: float = 0.1  # 10% tolerance

    # Scoring
    critical_issue_weight: float = 0.15
    warning_weight: float = 0.05
    min_score: float = 0.0

    @classmethod
    def from_env(cls) -> "ValidationConfig":
        """Create config from environment variables."""
        return cls(
            enable_docx=os.getenv("OPENCLAW_VALIDATE_DOCX", "1").strip() not in {"0", "false", "no"},
            enable_xlsx=os.getenv("OPENCLAW_VALIDATE_XLSX", "1").strip() not in {"0", "false", "no"},
            docx_check_fonts=os.getenv("OPENCLAW_VALIDATE_FONTS", "1").strip() not in {"0", "false", "no"},
            docx_check_styles=os.getenv("OPENCLAW_VALIDATE_STYLES", "1").strip() not in {"0", "false", "no"},
            docx_check_tables=os.getenv("OPENCLAW_VALIDATE_TABLES", "1").strip() not in {"0", "false", "no"},
            docx_font_size_delta=float(os.getenv("OPENCLAW_VALIDATE_FONT_DELTA", "0.5")),
            xlsx_check_fonts=os.getenv("OPENCLAW_VALIDATE_CELL_FONTS", "1").strip() not in {"0", "false", "no"},
            xlsx_check_borders=os.getenv("OPENCLAW_VALIDATE_BORDERS", "1").strip() not in {"0", "false", "no"},
            xlsx_check_fills=os.getenv("OPENCLAW_VALIDATE_FILLS", "1").strip() not in {"0", "false", "no"},
            xlsx_check_merged=os.getenv("OPENCLAW_VALIDATE_MERGED_CELLS", "1").strip() not in {"0", "false", "no"},
            xlsx_dimension_delta=float(os.getenv("OPENCLAW_VALIDATE_DIMENSION_DELTA", "0.1")),
        )


def _font_repr(run: Any) -> str:
    """Get string representation of font properties."""
    try:
        parts = []
        if run.font.name:
            parts.append(f"name={run.font.name}")
        if run.font.size:
            parts.append(f"size={run.font.size.pt}pt")
        if run.font.bold:
            parts.append("bold")
        if run.font.italic:
            parts.append("italic")
        if run.font.underline:
            parts.append("underline")
        return " ".join(parts) if parts else "default"
    except Exception:
        return "unknown"


def _paragraph_style_repr(para: Paragraph) -> str:
    """Get string representation of paragraph style."""
    try:
        if para.style and para.style.name:
            return para.style.name
        return "Normal"
    except Exception:
        return "unknown"


def _border_repr(border: Any) -> str:
    """Get string representation of border."""
    try:
        if not border or not border.style:
            return "none"
        return f"{border.style.name} {border.width.pt if hasattr(border.width, 'pt') else '?'}pt"
    except Exception:
        return "unknown"


def _fill_repr(fill: Any) -> str:
    """Get string representation of cell fill."""
    try:
        if not fill or not fill.start_color:
            return "none"
        if fill.start_color.type == "rgb":
            rgb = fill.start_color.rgb
            if rgb and hasattr(rgb, "hex"):
                return f"#{rgb.hex}"
        return "pattern"
    except Exception:
        return "unknown"


def _alignment_repr(align: Any) -> str:
    """Get string representation of cell alignment."""
    try:
        if not align:
            return "default"
        parts = []
        if align.horizontal:
            parts.append(f"h={align.horizontal}")
        if align.vertical:
            parts.append(f"v={align.vertical}")
        if align.wrap_text:
            parts.append("wrap")
        return " ".join(parts) if parts else "default"
    except Exception:
        return "unknown"


class DocxStructureValidator:
    """Validates DOCX structure preservation."""

    def __init__(self, config: ValidationConfig | None = None):
        self.config = config or ValidationConfig()

    def validate(
        self,
        original_path: Path,
        translated_path: Path,
    ) -> ValidationResult:
        """Compare original and translated DOCX at structure level."""
        if not DOCX_AVAILABLE:
            return ValidationResult(
                file_name=translated_path.name,
                file_path=str(translated_path),
                format_type="docx",
                valid=False,
                issues=[],
            )

        result = ValidationResult(
            file_name=translated_path.name,
            file_path=str(translated_path),
            format_type="docx",
            valid=True,
        )

        try:
            orig_doc = Document(str(original_path))
            trans_doc = Document(str(translated_path))
        except Exception as e:
            result.add_issue(ValidationIssue(
                category=Category.STYLE,
                severity=Severity.CRITICAL,
                location="file",
                element_type="document",
                expected="valid DOCX file",
                actual=f"read error: {e}",
                hint="Ensure translated file is a valid DOCX document",
            ))
            result.calculate_score()
            return result

        # Validate paragraphs
        if self.config.docx_check_fonts or self.config.docx_check_styles:
            self._compare_paragraphs(orig_doc, trans_doc, result)

        # Validate tables
        if self.config.docx_check_tables:
            self._compare_tables(orig_doc, trans_doc, result)

        # Validate headers/footers
        if self.config.docx_check_headers_footers:
            self._compare_headers_footers(orig_doc, trans_doc, result)

        result.calculate_score(
            critical_weight=self.config.critical_issue_weight,
            warning_weight=self.config.warning_weight,
        )
        return result

    def _compare_paragraphs(
        self,
        orig_doc: Document,
        trans_doc: Document,
        result: ValidationResult,
    ) -> None:
        """Compare paragraph styles and fonts."""
        orig_paras = [p for p in orig_doc.paragraphs if p.text.strip()]
        trans_paras = [p for p in trans_doc.paragraphs if p.text.strip()]

        # Compare paragraph count
        if len(orig_paras) != len(trans_paras):
            result.add_issue(ValidationIssue(
                category=Category.PARAGRAPH,
                severity=Severity.WARNING,
                location="document",
                element_type="paragraph_count",
                expected=f"{len(orig_paras)} paragraphs",
                actual=f"{len(trans_paras)} paragraphs",
                hint=f"Original has {len(orig_paras)} paragraphs, translated has {len(trans_paras)}",
            ))

        # Compare style and font for each paragraph
        for i, (orig, trans) in enumerate(zip(orig_paras, trans_paras), start=1):
            location = f"p:{i}"

            # Check style
            if self.config.docx_check_styles:
                orig_style = _paragraph_style_repr(orig)
                trans_style = _paragraph_style_repr(trans)
                if orig_style != trans_style:
                    result.add_issue(ValidationIssue(
                        category=Category.STYLE,
                        severity=Severity.WARNING,
                        location=location,
                        element_type="paragraph",
                        expected=f'style="{orig_style}"',
                        actual=f'style="{trans_style}"',
                        hint=f'Apply style "{orig_style}" to paragraph {i}',
                    ))

            # Check fonts in runs
            if self.config.docx_check_fonts:
                orig_runs = list(orig.runs)
                trans_runs = list(trans.runs)

                for j, (orig_run, trans_run) in enumerate(zip(orig_runs, trans_runs)):
                    run_location = f"{location}:run{j+1}"

                    # Compare font name
                    if orig_run.font.name and trans_run.font.name:
                        if orig_run.font.name != trans_run.font.name:
                            result.add_issue(ValidationIssue(
                                category=Category.FONT,
                                severity=Severity.WARNING,
                                location=run_location,
                                element_type="run",
                                expected=f'font="{orig_run.font.name}"',
                                actual=f'font="{trans_run.font.name}"',
                                hint=f'Set font to {orig_run.font.name} for run {j+1} in paragraph {i}',
                            ))

                    # Compare font size (with delta tolerance)
                    if orig_run.font.size and trans_run.font.size:
                        orig_size = orig_run.font.size.pt
                        trans_size = trans_run.font.size.pt
                        if abs(orig_size - trans_size) > self.config.docx_font_size_delta:
                            result.add_issue(ValidationIssue(
                                category=Category.FONT,
                                severity=Severity.INFO,
                                location=run_location,
                                element_type="run",
                                expected=f"size={orig_size}pt",
                                actual=f"size={trans_size}pt",
                                hint=f"Adjust font size to {orig_size}pt (within {self.config.docx_font_size_delta}pt tolerance)",
                            ))
                else:
                    # Count passes for runs that were compared
                    if orig_runs:
                        result.add_pass()

    def _compare_tables(
        self,
        orig_doc: Document,
        trans_doc: Document,
        result: ValidationResult,
    ) -> None:
        """Compare table structure."""
        orig_tables = list(orig_doc.tables)
        trans_tables = list(trans_doc.tables)

        if len(orig_tables) != len(trans_tables):
            result.add_issue(ValidationIssue(
                category=Category.TABLE,
                severity=Severity.CRITICAL,
                location="document",
                element_type="table_count",
                expected=f"{len(orig_tables)} tables",
                actual=f"{len(trans_tables)} tables",
                hint=f"Original has {len(orig_tables)} tables, translated has {len(trans_tables)}",
            ))

        for i, (orig, trans) in enumerate(zip(orig_tables, trans_tables), start=1):
            table_location = f"table:{i}"

            # Check dimensions
            orig_rows = len(orig.rows)
            orig_cols = len(orig.columns)
            trans_rows = len(trans.rows)
            trans_cols = len(trans.columns)

            if orig_rows != trans_rows or orig_cols != trans_cols:
                result.add_issue(ValidationIssue(
                    category=Category.TABLE,
                    severity=Severity.CRITICAL,
                    location=table_location,
                    element_type="table_dimensions",
                    expected=f"{orig_rows}x{orig_cols}",
                    actual=f"{trans_rows}x{trans_cols}",
                    hint=f"Restore table {i} to {orig_rows} rows x {orig_cols} columns",
                ))

            # Check merged cells and borders
            self._compare_table_cells(orig, trans, table_location, result)

    def _compare_table_cells(
        self,
        orig: Table,
        trans: Table,
        table_location: str,
        result: ValidationResult,
    ) -> None:
        """Compare table cell borders and merges."""
        try:
            # Check merged regions by comparing cell spans
            for r_idx, (orig_row, trans_row) in enumerate(zip(orig.rows, trans.rows), start=1):
                for c_idx, (orig_cell, trans_cell) in enumerate(zip(orig_row.cells, trans_row.cells), start=1):
                    cell_location = f"{table_location}:r{r_idx}:c{c_idx}"

                    # Check for merged cells (spanned cells have same _tc element)
                    orig_tc = orig_cell._tc
                    trans_tc = trans_cell._tc

                    # Count how many cells reference the same tc element
                    orig_span_count = sum(
                        1 for r in orig.rows
                        for c in r.cells
                        if c._tc == orig_tc
                    )
                    trans_span_count = sum(
                        1 for r in trans.rows
                        for c in r.cells
                        if c._tc == trans_tc
                    )

                    if orig_span_count != trans_span_count:
                        result.add_issue(ValidationIssue(
                            category=Category.TABLE_BORDERS,
                            severity=Severity.CRITICAL,
                            location=cell_location,
                            element_type="merged_cell",
                            expected=f"merge span={orig_span_count}",
                            actual=f"merge span={trans_span_count}",
                            hint=f"Check merge status for cell {cell_location} (original spans {orig_span_count} cells)",
                        ))

                    # Check borders on first cell edge
                    if r_idx == 1 and c_idx == 1:
                        try:
                            orig_borders = orig_cell._tc.tcPr.border if orig_cell._tc.tcPr else None
                            trans_borders = trans_cell._tc.tcPr.border if trans_cell._tc.tcPr else None
                            # Basic check - full border comparison would require deeper XML inspection
                            if (orig_borders is None) != (trans_borders is None):
                                result.add_issue(ValidationIssue(
                                    category=Category.TABLE_BORDERS,
                                    severity=Severity.WARNING,
                                    location=cell_location,
                                    element_type="cell_borders",
                                    expected="borders present" if orig_borders else "no borders",
                                    actual="borders present" if trans_borders else "no borders",
                                    hint=f"Restore table borders to match original for {cell_location}",
                                ))
                        except Exception:
                            pass

        except Exception as e:
            result.add_issue(ValidationIssue(
                category=Category.TABLE,
                severity=Severity.WARNING,
                location=table_location,
                element_type="table_structure",
                expected="valid table structure",
                actual=f"comparison error: {e}",
                hint="Manual review of table structure recommended",
            ))

    def _compare_headers_footers(
        self,
        orig_doc: Document,
        trans_doc: Document,
        result: ValidationResult,
    ) -> None:
        """Compare header/footer presence."""
        try:
            orig_sections = list(orig_doc.sections)
            trans_sections = list(trans_doc.sections)

            if len(orig_sections) != len(trans_sections):
                result.add_issue(ValidationIssue(
                    category=Category.HEADER_FOOTER,
                    severity=Severity.INFO,
                    location="document",
                    element_type="section_count",
                    expected=f"{len(orig_sections)} sections",
                    actual=f"{len(trans_sections)} sections",
                    hint="Section count differs between documents",
                ))

            for i, (orig, trans) in enumerate(zip(orig_sections, trans_sections), start=1):
                # Check headers
                orig_header = orig.header
                trans_header = trans.header
                if (orig_header is not None) != (trans_header is not None):
                    result.add_issue(ValidationIssue(
                        category=Category.HEADER_FOOTER,
                        severity=Severity.WARNING,
                        location=f"section:{i}",
                        element_type="header",
                        expected="header present" if orig_header else "no header",
                        actual="header present" if trans_header else "no header",
                        hint=f"Restore header in section {i} to match original",
                    ))

                # Check footers
                orig_footer = orig.footer
                trans_footer = trans.footer
                if (orig_footer is not None) != (trans_footer is not None):
                    result.add_issue(ValidationIssue(
                        category=Category.HEADER_FOOTER,
                        severity=Severity.WARNING,
                        location=f"section:{i}",
                        element_type="footer",
                        expected="footer present" if orig_footer else "no footer",
                        actual="footer present" if trans_footer else "no footer",
                        hint=f"Restore footer in section {i} to match original",
                    ))

        except Exception as e:
            result.add_issue(ValidationIssue(
                category=Category.HEADER_FOOTER,
                severity=Severity.INFO,
                location="document",
                element_type="header_footer_check",
                expected="valid sections",
                actual=f"comparison error: {e}",
                hint="Manual review of headers/footers recommended",
            ))


class XlsxStructureValidator:
    """Validates XLSX structure preservation."""

    def __init__(self, config: ValidationConfig | None = None):
        self.config = config or ValidationConfig()

    def validate(
        self,
        original_path: Path,
        translated_path: Path,
    ) -> ValidationResult:
        """Compare original and translated XLSX at structure level."""
        if not OPENPYXL_AVAILABLE:
            return ValidationResult(
                file_name=translated_path.name,
                file_path=str(translated_path),
                format_type="xlsx",
                valid=False,
                issues=[],
            )

        result = ValidationResult(
            file_name=translated_path.name,
            file_path=str(translated_path),
            format_type="xlsx",
            valid=True,
        )

        try:
            orig_wb = openpyxl.load_workbook(str(original_path), data_only=False)
            trans_wb = openpyxl.load_workbook(str(translated_path), data_only=False)
        except Exception as e:
            result.add_issue(ValidationIssue(
                category=Category.CELL_FONT,
                severity=Severity.CRITICAL,
                location="file",
                element_type="workbook",
                expected="valid XLSX file",
                actual=f"read error: {e}",
                hint="Ensure translated file is a valid XLSX workbook",
            ))
            result.calculate_score()
            return result

        try:
            # Compare worksheets
            self._compare_worksheets(orig_wb, trans_wb, result)

            # Compare each sheet's cells
            for sheet_name in orig_wb.sheetnames:
                if sheet_name in trans_wb.sheetnames:
                    orig_ws = orig_wb[sheet_name]
                    trans_ws = trans_wb[sheet_name]
                    self._compare_sheet_cells(sheet_name, orig_ws, trans_ws, result)

        finally:
            orig_wb.close()
            trans_wb.close()

        result.calculate_score(
            critical_weight=self.config.critical_issue_weight,
            warning_weight=self.config.warning_weight,
        )
        return result

    def _compare_worksheets(
        self,
        orig_wb: openpyxl.Workbook,
        trans_wb: openpyxl.Workbook,
        result: ValidationResult,
    ) -> None:
        """Compare worksheet structure."""
        orig_sheets = set(orig_wb.sheetnames)
        trans_sheets = set(trans_wb.sheetnames)

        if orig_sheets != trans_sheets:
            missing = orig_sheets - trans_sheets
            extra = trans_sheets - orig_sheets

            if missing:
                result.add_issue(ValidationIssue(
                    category=Category.CELL_FONT,
                    severity=Severity.CRITICAL,
                    location="workbook",
                    element_type="worksheet_names",
                    expected=f"sheets: {sorted(orig_sheets)}",
                    actual=f"sheets: {sorted(trans_sheets)}",
                    hint=f"Missing worksheets: {missing}",
                ))

            if extra:
                result.add_issue(ValidationIssue(
                    category=Category.CELL_FONT,
                    severity=Severity.INFO,
                    location="workbook",
                    element_type="worksheet_names",
                    expected=f"sheets: {sorted(orig_sheets)}",
                    actual=f"sheets: {sorted(trans_sheets)} (extra: {extra})",
                    hint=f"Extra worksheets found: {extra}",
                ))

    def _compare_sheet_cells(
        self,
        sheet_name: str,
        orig_ws: openpyxl.worksheet.worksheet.Worksheet,
        trans_ws: openpyxl.worksheet.worksheet.Worksheet,
        result: ValidationResult,
    ) -> None:
        """Compare cells in a worksheet."""
        # Determine comparison range
        orig_max_row = max(orig_ws.max_row, 1)
        orig_max_col = max(orig_ws.max_column, 1)
        trans_max_row = max(trans_ws.max_row, 1)
        trans_max_col = max(trans_ws.max_column, 1)

        max_row = max(orig_max_row, trans_max_row)
        max_col = max(orig_max_col, trans_max_col)

        # Check dimensions
        if orig_max_row != trans_max_row or orig_max_col != trans_max_col:
            result.add_issue(ValidationIssue(
                category=Category.DIMENSIONS,
                severity=Severity.WARNING,
                location=f"{sheet_name}!dimensions",
                element_type="used_range",
                expected=f"{orig_max_row} rows x {orig_max_col} cols",
                actual=f"{trans_max_row} rows x {trans_max_col} cols",
                hint=f"Sheet used range differs: expected {orig_max_row}x{orig_max_col}, got {trans_max_row}x{trans_max_col}",
            ))

        # Compare merged regions first
        if self.config.xlsx_check_merged:
            self._compare_merged_regions(sheet_name, orig_ws, trans_ws, result)

        # Compare cell properties
        for row_idx in range(1, max_row + 1):
            for col_idx in range(1, max_col + 1):
                orig_cell = orig_ws.cell(row=row_idx, column=col_idx)
                trans_cell = trans_ws.cell(row=row_idx, column=col_idx)
                cell_addr = trans_cell.coordinate
                location = f"{sheet_name}!{cell_addr}"

                # Skip if both cells are effectively empty
                if not orig_cell.value and not trans_cell.value:
                    continue

                self._compare_cell(location, orig_cell, trans_cell, result)

        # Compare column widths and row heights
        if self.config.xlsx_check_dimensions:
            self._compare_dimensions(sheet_name, orig_ws, trans_ws, result)

    def _compare_cell(
        self,
        location: str,
        orig_cell: openpyxl.cell.cell.Cell,
        trans_cell: openpyxl.cell.cell.Cell,
        result: ValidationResult,
    ) -> None:
        """Compare individual cell properties."""
        # Skip formula cells (values should match, format less critical)
        if orig_cell.data_type == "f":
            return

        # Compare fonts
        if self.config.xlsx_check_fonts:
            orig_font = orig_cell.font
            trans_font = trans_cell.font

            if orig_font and trans_font:
                issues = []

                if orig_font.name and trans_font.name and orig_font.name != trans_font.name:
                    issues.append(f"name: {orig_font.name} -> {trans_font.name}")
                if orig_font.size and trans_font.size and abs(orig_font.size - trans_font.size) > 0.5:
                    issues.append(f"size: {orig_font.size}pt -> {trans_font.size}pt")
                if orig_font.bold != trans_font.bold:
                    issues.append(f"bold: {orig_font.bold} -> {trans_font.bold}")
                if orig_font.italic != trans_font.italic:
                    issues.append(f"italic: {orig_font.italic} -> {trans_font.italic}")

                if issues:
                    result.add_issue(ValidationIssue(
                        category=Category.CELL_FONT,
                        severity=Severity.INFO,
                        location=location,
                        element_type="cell",
                        expected=f"font: {orig_font.name or 'default'} {orig_font.size or '?'}pt {'bold' if orig_font.bold else ''} {'italic' if orig_font.italic else ''}",
                        actual=f"font: {trans_font.name or 'default'} {trans_font.size or '?'}pt {'bold' if trans_font.bold else ''} {'italic' if trans_font.italic else ''}",
                        hint=f"Cell {location}: " + "; ".join(issues),
                    ))
            else:
                result.add_pass()

        # Compare fills
        if self.config.xlsx_check_fills:
            orig_fill = orig_cell.fill
            trans_fill = trans_cell.fill

            if orig_fill and trans_fill:
                orig_repr = _fill_repr(orig_fill)
                trans_repr = _fill_repr(trans_fill)

                if orig_repr != trans_repr and orig_repr != "none" and trans_repr != "none":
                    result.add_issue(ValidationIssue(
                        category=Category.CELL_FILL,
                        severity=Severity.INFO,
                        location=location,
                        element_type="cell",
                        expected=f"fill: {orig_repr}",
                        actual=f"fill: {trans_repr}",
                        hint=f"Restore fill color to {orig_repr} for cell {location}",
                    ))
            else:
                result.add_pass()

        # Compare borders
        if self.config.xlsx_check_borders:
            orig_border = orig_cell.border
            trans_border = trans_cell.border

            border_issues = []
            for side in ["left", "right", "top", "bottom"]:
                orig_side = getattr(orig_border, side)
                trans_side = getattr(trans_border, side)
                orig_repr = _border_repr(orig_side)
                trans_repr = _border_repr(trans_side)

                if orig_repr != trans_repr and orig_repr != "none":
                    border_issues.append(f"{side}: {orig_repr} -> {trans_repr}")

            if border_issues:
                result.add_issue(ValidationIssue(
                    category=Category.CELL_BORDER,
                    severity=Severity.INFO,
                    location=location,
                    element_type="cell",
                    expected=f"borders match",
                    actual=f"border differences: " + "; ".join(border_issues),
                    hint=f"Restore borders for cell {location}: " + "; ".join(border_issues),
                ))
            else:
                result.add_pass()

        # Compare alignment
        if self.config.xlsx_check_alignment:
            orig_align = orig_cell.alignment
            trans_align = trans_cell.alignment

            if orig_align and trans_align:
                orig_repr = _alignment_repr(orig_align)
                trans_repr = _alignment_repr(trans_align)

                if orig_repr != trans_repr and orig_repr != "default":
                    result.add_issue(ValidationIssue(
                        category=Category.CELL_ALIGNMENT,
                        severity=Severity.INFO,
                        location=location,
                        element_type="cell",
                        expected=f"alignment: {orig_repr}",
                        actual=f"alignment: {trans_repr}",
                        hint=f"Restore alignment to {orig_repr} for cell {location}",
                    ))
            else:
                result.add_pass()

    def _compare_merged_regions(
        self,
        sheet_name: str,
        orig_ws: openpyxl.worksheet.worksheet.Worksheet,
        trans_ws: openpyxl.worksheet.worksheet.Worksheet,
        result: ValidationResult,
    ) -> None:
        """Compare merged cell regions."""
        orig_merged = {str(rng) for rng in orig_ws.merged_cells.ranges}
        trans_merged = {str(rng) for rng in trans_ws.merged_cells.ranges}

        if orig_merged != trans_merged:
            missing = orig_merged - trans_merged
            extra = trans_merged - orig_merged

            if missing:
                for rng in sorted(missing):
                    result.add_issue(ValidationIssue(
                        category=Category.MERGED_REGIONS,
                        severity=Severity.CRITICAL,
                        location=f"{sheet_name}!{rng}",
                        element_type="merged_cells",
                        expected=f"merged: {rng}",
                        actual="not merged",
                        hint=f"Restore merged region {rng} in sheet {sheet_name}",
                    ))

            if extra:
                for rng in sorted(extra):
                    result.add_issue(ValidationIssue(
                        category=Category.MERGED_REGIONS,
                        severity=Severity.WARNING,
                        location=f"{sheet_name}!{rng}",
                        element_type="merged_cells",
                        expected="not merged",
                        actual=f"merged: {rng}",
                        hint=f"Unexpected merged region {rng} in translated sheet",
                    ))

    def _compare_dimensions(
        self,
        sheet_name: str,
        orig_ws: openpyxl.worksheet.worksheet.Worksheet,
        trans_ws: openpyxl.worksheet.worksheet.Worksheet,
        result: ValidationResult,
    ) -> None:
        """Compare column widths and row heights."""
        delta = self.config.xlsx_dimension_delta

        # Compare column widths
        for col_letter, orig_dim in orig_ws.column_dimensions.items():
            trans_dim = trans_ws.column_dimensions.get(col_letter)
            if orig_dim.width and trans_dim and trans_dim.width:
                orig_w = float(orig_dim.width)
                trans_w = float(trans_dim.width)
                if abs(orig_w - trans_w) / max(orig_w, 0.1) > delta:
                    result.add_issue(ValidationIssue(
                        category=Category.DIMENSIONS,
                        severity=Severity.INFO,
                        location=f"{sheet_name}!{col_letter}",
                        element_type="column_width",
                        expected=f"width: {orig_w:.2f}",
                        actual=f"width: {trans_w:.2f}",
                        hint=f"Column {col_letter} width differs (expected {orig_w:.2f}, got {trans_w:.2f})",
                    ))

        # Compare row heights
        for row_idx, orig_dim in orig_ws.row_dimensions.items():
            trans_dim = trans_ws.row_dimensions.get(row_idx)
            if orig_dim.height and trans_dim and trans_dim.height:
                orig_h = float(orig_dim.height)
                trans_h = float(trans_dim.height)
                if abs(orig_h - trans_h) / max(orig_h, 0.1) > delta:
                    result.add_issue(ValidationIssue(
                        category=Category.DIMENSIONS,
                        severity=Severity.INFO,
                        location=f"{sheet_name}!{row_idx}",
                        element_type="row_height",
                        expected=f"height: {orig_h:.2f}",
                        actual=f"height: {trans_h:.2f}",
                        hint=f"Row {row_idx} height differs (expected {orig_h:.2f}, got {trans_h:.2f})",
                    ))


class ValidationReportGenerator:
    """Generates reports from validation results."""

    def generate_markdown(
        self,
        results: list[ValidationResult],
        *,
        include_hints: bool = True,
        max_issues_per_category: int = 20,
    ) -> str:
        """Generate LLM-consumable markdown report."""
        lines = ["# Detail Validation Report", ""]

        # Summary section
        total_checks = sum(r.total_checks for r in results)
        total_passed = sum(r.passed for r in results)
        total_warnings = sum(r.warnings for r in results)
        total_failed = sum(r.failed for r in results)
        overall_score = sum(r.format_fidelity_score for r in results) / max(len(results), 1)

        lines.extend([
            "## Summary",
            f"**Total Checks:** {total_checks}",
            f"**Passed:** {total_passed}",
            f"**Warnings:** {total_warnings}",
            f"**Failed:** {total_failed}",
            f"**Overall Score:** {overall_score:.1%}",
            "",
        ])

        # Issues by file
        if any(r.issues for r in results):
            lines.append("## Issues Found")

            for result in results:
                if not result.issues:
                    continue

                lines.extend([
                    f"",
                    f"### {result.file_name}",
                    f"**Score:** {result.format_fidelity_score:.1%} | "
                    f"**Issues:** {result.total_checks} ({result.failed} failed, {result.warnings} warnings)",
                    ""
                ])

                # Group by severity
                critical = [i for i in result.issues if i.severity == Severity.CRITICAL]
                warnings = [i for i in result.issues if i.severity == Severity.WARNING]
                info = [i for i in result.issues if i.severity == Severity.INFO]

                if critical:
                    lines.append("#### Errors")
                    for issue in critical[:max_issues_per_category]:
                        lines.extend(self._format_issue(issue, include_hints))
                    lines.append("")

                if warnings:
                    lines.append("#### Warnings")
                    for issue in warnings[:max_issues_per_category]:
                        lines.extend(self._format_issue(issue, include_hints))
                    lines.append("")

                if info and include_hints:
                    lines.append("#### Info")
                    for issue in info[:max_issues_per_category]:
                        lines.extend(self._format_issue(issue, include_hints))
                    lines.append("")

        # Category summary
        lines.extend(["", "## Summary by Category", ""])
        category_totals: dict[str, dict[str, int]] = {}

        for result in results:
            for issue in result.issues:
                if issue.category not in category_totals:
                    category_totals[issue.category] = {"failed": 0, "warning": 0, "info": 0}
                if issue.severity == Severity.CRITICAL:
                    category_totals[issue.category]["failed"] += 1
                elif issue.severity == Severity.WARNING:
                    category_totals[issue.category]["warning"] += 1
                else:
                    category_totals[issue.category]["info"] += 1

        if category_totals:
            lines.append("| Category | Failed | Warnings | Info |")
            lines.append("|----------|--------|----------|------|")
            for cat in sorted(category_totals.keys()):
                counts = category_totals[cat]
                lines.append(f"| {cat} | {counts['failed']} | {counts['warning']} | {counts['info']} |")

        # Auto-fix recommendations
        if include_hints and any(r.issues for r in results):
            lines.extend(["", "## Recommendations for LLM Fix", ""])
            hint_counter = 1
            for result in results:
                for issue in result.issues:
                    if issue.hint and issue.severity in (Severity.CRITICAL, Severity.WARNING):
                        lines.append(f"{hint_counter}. {issue.hint}")
                        hint_counter += 1

        return "\n".join(lines)

    def _format_issue(self, issue: ValidationIssue, include_hints: bool) -> list[str]:
        """Format a single issue as markdown."""
        severity_icon = {
            Severity.CRITICAL: "",
            Severity.WARNING: "⚠️",
            Severity.INFO: "",
        }

        lines = [
            f"- {severity_icon.get(issue.severity, '')} **[{issue.category.upper()}]** {issue.location}",
            f"  - **Problem:** {issue.element_type}",
            f"  - **Expected:** `{issue.expected}`",
            f"  - **Actual:** `{issue.actual}`",
        ]

        if include_hints and issue.hint:
            lines.append(f"  - **Fix:** {issue.hint}")

        return lines

    def generate_summary(self, results: list[ValidationResult]) -> dict[str, Any]:
        """Generate summary for quality gate integration."""
        total_checks = sum(r.total_checks for r in results)
        total_passed = sum(r.passed for r in results)
        total_warnings = sum(r.warnings for r in results)
        total_failed = sum(r.failed for r in results)
        overall_score = sum(r.format_fidelity_score for r in results) / max(len(results), 1)

        # Group by category
        by_category: dict[str, dict[str, int]] = {}
        for result in results:
            for issue in result.issues:
                if issue.category not in by_category:
                    by_category[issue.category] = {"passed": 0, "failed": 0}
                if issue.severity == Severity.CRITICAL:
                    by_category[issue.category]["failed"] += 1
                else:
                    by_category[issue.category]["passed"] += 1

        return {
            "total_checks": total_checks,
            "passed": total_passed,
            "warnings": total_warnings,
            "failed": total_failed,
            "score": round(overall_score, 3),
            "by_category": by_category,
            "files_valid": sum(1 for r in results if r.valid),
            "total_files": len(results),
        }

    def extract_fix_hints(self, results: list[ValidationResult]) -> list[str]:
        """Extract fix hints for auto-fix feedback."""
        hints = []
        for result in results:
            for issue in result.issues:
                if issue.hint and issue.severity in (Severity.CRITICAL, Severity.WARNING):
                    hints.append(f"{result.file_name}: {issue.hint}")
        return hints


def validate_file_pair(
    original_path: Path,
    translated_path: Path,
    *,
    config: ValidationConfig | None = None,
) -> ValidationResult:
    """Validate a single original/translated file pair.

    Args:
        original_path: Path to original/source document
        translated_path: Path to translated document
        config: Validation configuration (uses defaults if None)

    Returns:
        ValidationResult with all issues found

    Raises:
        ValueError: If file formats don't match or are unsupported
        FileNotFoundError: If either file doesn't exist
    """
    original_path = Path(original_path).expanduser().resolve()
    translated_path = Path(translated_path).expanduser().resolve()

    if not original_path.exists():
        raise FileNotFoundError(f"Original file not found: {original_path}")
    if not translated_path.exists():
        raise FileNotFoundError(f"Translated file not found: {translated_path}")

    orig_ext = original_path.suffix.lower()
    trans_ext = translated_path.suffix.lower()

    if orig_ext != trans_ext:
        raise ValueError(f"File format mismatch: {orig_ext} vs {trans_ext}")

    config = config or ValidationConfig.from_env()

    if orig_ext == ".docx" and config.enable_docx:
        validator = DocxStructureValidator(config)
    elif orig_ext == ".xlsx" and config.enable_xlsx:
        validator = XlsxStructureValidator(config)
    else:
        raise ValueError(f"Unsupported file format: {orig_ext}")

    return validator.validate(original_path, translated_path)


def validate_job_artifacts(
    review_dir: Path,
    original_files: list[Path],
    translated_files: list[Path],
    *,
    config: ValidationConfig | None = None,
) -> dict[str, ValidationResult]:
    """Validate all file pairs for a job.

    Args:
        review_dir: Path to review directory
        original_files: List of original file paths
        translated_files: List of translated file paths
        config: Validation configuration

    Returns:
        Dict mapping file_name -> ValidationResult
    """
    review_dir = Path(review_dir).expanduser().resolve()
    config = config or ValidationConfig.from_env()
    results: dict[str, ValidationResult] = {}

    # Simple matching by filename (can be enhanced with task_bundle_builder logic)
    orig_by_name = {Path(f).name: f for f in original_files}
    trans_by_name = {Path(f).name: f for f in translated_files}

    for name, orig_path in orig_by_name.items():
        # Find corresponding translated file
        trans_path = trans_by_name.get(name)
        if not trans_path:
            # Try to find a file with same base but different extension pattern
            for trans_name, trans_file in trans_by_name.items():
                if Path(orig_path).stem == Path(trans_file).stem:
                    trans_path = trans_file
                    break

        if trans_path:
            try:
                result = validate_file_pair(orig_path, trans_path, config=config)
                results[name] = result
            except Exception as e:
                # Create error result
                results[name] = ValidationResult(
                    file_name=name,
                    file_path=str(trans_path),
                    format_type=Path(orig_path).suffix[1:],
                    valid=False,
                    issues=[],
                )
                results[name].add_issue(ValidationIssue(
                    category="validation_error",
                    severity=Severity.CRITICAL,
                    location="file",
                    element_type="validation",
                    expected="successful validation",
                    actual=f"error: {e}",
                    hint=f"Validation failed for {name}: {e}",
                ))

    return results


def main() -> int:
    """CLI entry point for detail validator."""
    import argparse

    parser = argparse.ArgumentParser(
        description="Validate format preservation between original and translated documents"
    )
    parser.add_argument("--original", required=True, help="Path to original document")
    parser.add_argument("--translated", required=True, help="Path to translated document")
    parser.add_argument("--output", help="Output JSON file path")
    parser.add_argument("--report", help="Output markdown report path")
    parser.add_argument("--format", choices=["json", "markdown", "both"], default="both")
    args = parser.parse_args()

    try:
        result = validate_file_pair(
            Path(args.original),
            Path(args.translated),
        )
    except Exception as e:
        print(f"Error: {e}", file=__import__("sys").stderr)
        return 1

    generator = ValidationReportGenerator()

    if args.format in ("json", "both"):
        output = {
            "summary": generator.generate_summary([result]),
            "files": [result.to_dict()],
            "visual_report": generator.generate_markdown([result]),
        }

        if args.output:
            Path(args.output).parent.mkdir(parents=True, exist_ok=True)
            Path(args.output).write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
        else:
            print(json.dumps(output, ensure_ascii=False, indent=2))

    if args.format in ("markdown", "both"):
        report = generator.generate_markdown([result])
        if args.report:
            Path(args.report).parent.mkdir(parents=True, exist_ok=True)
            Path(args.report).write_text(report, encoding="utf-8")
        elif args.format == "markdown":
            print(report)

    return 0


if __name__ == "__main__":
    import sys
    sys.exit(main())
