#!/usr/bin/env python3
"""Knowledge-base indexing and retrieval for OpenClaw V4.1."""

from __future__ import annotations

import csv
import json
import re
import sqlite3
import subprocess
import sys
from pathlib import Path
from typing import Any

from docx import Document

from scripts.skill_clawrag_bridge import clawrag_search, clawrag_sync
from scripts.v4_runtime import (
    KB_SUPPORTED_EXTENSIONS,
    SOURCE_GROUP_WEIGHTS,
    compute_sha256,
    infer_source_group,
    json_dumps,
    utc_now_iso,
)

_KB_FTS_AVAILABLE: bool | None = None


def _reference_like_filters(*, kb_root: Path, kb_company: str) -> tuple[str, str]:
    ref_root = (kb_root / "30_Reference").expanduser().resolve()
    company_root = (ref_root / kb_company).expanduser().resolve()
    return f"{ref_root}/%", f"{company_root}/%"


def _allow_kb_path(
    path: str,
    *,
    kb_root: Path | None,
    kb_company: str,
    isolation_mode: str,
) -> bool:
    p = str(path or "").strip()
    if not p:
        return False

    if kb_root:
        kb_root_abs = str(kb_root.expanduser().resolve())
        if not (p.startswith(kb_root_abs + "/") or p == kb_root_abs):
            return False
    try:
        if not Path(p).expanduser().exists():
            return False
    except OSError:
        return False

    if not kb_root or not kb_company.strip():
        return True

    mode = (isolation_mode or "reference_only").strip().lower()
    if mode not in {"reference_only", "all"}:
        mode = "reference_only"

    ref_root = str((kb_root / "30_Reference").expanduser().resolve())
    company_root = str((kb_root / "30_Reference" / kb_company).expanduser().resolve())

    if mode == "all":
        return p.startswith(company_root + "/") or p == company_root

    # reference_only: only filter Reference; allow glossary/style/domain/templates globally
    if p.startswith(ref_root + "/"):
        return p.startswith(company_root + "/") or p == company_root
    return True


def _ensure_kb_fts(conn: sqlite3.Connection) -> bool:
    """Best-effort: enable FTS5 BM25 retrieval when supported by SQLite."""
    global _KB_FTS_AVAILABLE
    if _KB_FTS_AVAILABLE is False:
        return False

    row = conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name='kb_chunks_fts'"
    ).fetchone()
    if row:
        _KB_FTS_AVAILABLE = True
        return True

    try:
        conn.execute(
            "CREATE VIRTUAL TABLE kb_chunks_fts USING fts5(text, content='kb_chunks', content_rowid='id', tokenize='unicode61')"
        )
        conn.executescript(
            """
            CREATE TRIGGER IF NOT EXISTS kb_chunks_ai AFTER INSERT ON kb_chunks BEGIN
              INSERT INTO kb_chunks_fts(rowid, text) VALUES (new.id, new.text);
            END;
            CREATE TRIGGER IF NOT EXISTS kb_chunks_ad AFTER DELETE ON kb_chunks BEGIN
              INSERT INTO kb_chunks_fts(kb_chunks_fts, rowid, text) VALUES('delete', old.id, old.text);
            END;
            CREATE TRIGGER IF NOT EXISTS kb_chunks_au AFTER UPDATE ON kb_chunks BEGIN
              INSERT INTO kb_chunks_fts(kb_chunks_fts, rowid, text) VALUES('delete', old.id, old.text);
              INSERT INTO kb_chunks_fts(rowid, text) VALUES (new.id, new.text);
            END;
            """
        )
        conn.execute("INSERT INTO kb_chunks_fts(kb_chunks_fts) VALUES('rebuild')")
        conn.commit()
        _KB_FTS_AVAILABLE = True
        return True
    except sqlite3.OperationalError:
        _KB_FTS_AVAILABLE = False
        return False

try:  # Optional dependency
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional import
    PdfReader = None

try:  # Optional dependency
    from openpyxl import load_workbook
except Exception:  # pragma: no cover - optional import
    load_workbook = None


def _normalize_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _chunk_text(text: str, max_chars: int = 1200, overlap: int = 120) -> list[str]:
    norm = _normalize_text(text)
    if not norm:
        return []

    units = [u.strip() for u in re.split(r"(?:\n{2,}|(?<=[.!?])\s+)", norm) if u.strip()]
    chunks: list[str] = []
    cur = ""
    for unit in units:
        if not cur:
            cur = unit
            continue
        if len(cur) + 1 + len(unit) <= max_chars:
            cur = f"{cur} {unit}"
        else:
            chunks.append(cur)
            tail = cur[-overlap:] if overlap > 0 and len(cur) > overlap else cur
            cur = f"{tail} {unit}" if tail else unit
            if len(cur) > max_chars:
                chunks.append(cur[:max_chars])
                cur = cur[max_chars - overlap :] if overlap > 0 else ""
    if cur:
        chunks.append(cur)
    return [_normalize_text(c) for c in chunks if _normalize_text(c)]


def _extract_docx(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = _normalize_text(p.text)
        if t:
            lines.append(t)
    for t_idx, table in enumerate(doc.tables, start=1):
        lines.append(f"[Table {t_idx}]")
        for row in table.rows:
            cells = [_normalize_text(c.text).replace("\n", " / ") for c in row.cells]
            row_text = " | ".join([c for c in cells if c])
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def _extract_pdf(path: Path) -> str:
    pdftotext_bin = "/opt/homebrew/bin/pdftotext"
    if Path(pdftotext_bin).exists():
        try:
            proc = subprocess.run(
                [pdftotext_bin, "-layout", str(path), "-"],
                capture_output=True, text=True, timeout=60,
            )
            if proc.returncode == 0 and proc.stdout.strip():
                return _normalize_text(proc.stdout)
        except (subprocess.TimeoutExpired, OSError):
            pass
    # Fallback to pypdf
    if PdfReader is None:
        raise RuntimeError("Neither pdftotext nor pypdf available for .pdf")
    reader = PdfReader(str(path))
    lines: list[str] = []
    for page in reader.pages:
        txt = _normalize_text(page.extract_text() or "")
        if txt:
            lines.append(txt)
    return "\n".join(lines)


def _extract_text_plain(path: Path) -> str:
    return _normalize_text(path.read_text(encoding="utf-8", errors="ignore"))


def _extract_csv(path: Path) -> str:
    lines: list[str] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            cleaned = [_normalize_text(c) for c in row if _normalize_text(c)]
            if cleaned:
                lines.append(" | ".join(cleaned))
    return "\n".join(lines)


SHEETSMITH_SCRIPT = Path.home() / ".openclaw/workspace/skills/sheetsmith/scripts/sheetsmith.py"


def _extract_xlsx(path: Path) -> str:
    if SHEETSMITH_SCRIPT.exists():
        try:
            proc = subprocess.run(
                [sys.executable, str(SHEETSMITH_SCRIPT), "preview", str(path), "--rows", "9999"],
                capture_output=True, text=True, timeout=60,
            )
            if proc.returncode == 0 and proc.stdout.strip():
                return _normalize_text(proc.stdout)
        except (subprocess.TimeoutExpired, OSError):
            pass
    # Fallback to openpyxl
    if load_workbook is None:
        raise RuntimeError("Neither sheetsmith nor openpyxl available for .xlsx")
    wb = load_workbook(str(path), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"[Sheet] {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells: list[str] = []
            for value in row:
                if value is None:
                    continue
                text = _normalize_text(str(value))
                if text:
                    cells.append(text)
            if cells:
                lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines)


def extract_text(path: Path) -> tuple[str, str]:
    ext = path.suffix.lower()
    if ext == ".docx":
        return "docx", _extract_docx(path)
    if ext == ".pdf":
        return "pdf", _extract_pdf(path)
    if ext in {".md", ".txt"}:
        return ext[1:], _extract_text_plain(path)
    if ext == ".csv":
        return "csv", _extract_csv(path)
    if ext == ".xlsx":
        return "xlsx", _extract_xlsx(path)
    return "unknown", ""


def discover_kb_files(kb_root: Path) -> list[Path]:
    files: list[Path] = []
    for p in kb_root.rglob("*"):
        if not p.is_file():
            continue
        if p.name.startswith("~$") or p.name.startswith("."):
            continue
        low_path = str(p).lower().replace("\\", "/")
        # Do not index raw source uploads inside Reference projects by default.
        if "/30_reference/" in low_path and "/source/" in low_path:
            continue
        if p.suffix.lower() not in KB_SUPPORTED_EXTENSIONS:
            continue
        files.append(p)
    files.sort(key=lambda x: str(x).lower())
    return files


def sync_kb(
    *,
    conn: sqlite3.Connection,
    kb_root: Path,
    report_path: Path | None = None,
) -> dict[str, Any]:
    kb_root = kb_root.expanduser().resolve()
    files = discover_kb_files(kb_root)

    existing_rows = conn.execute("SELECT * FROM kb_files").fetchall()
    existing = {row["path"]: dict(row) for row in existing_rows}
    seen_paths = set()

    report: dict[str, Any] = {
        "ok": True,
        "kb_root": str(kb_root),
        "scanned_count": len(files),
        "created": 0,
        "updated": 0,
        "metadata_only": 0,
        "metadata_only_paths": [],
        "removed": 0,
        "skipped": 0,
        "errors": [],
        "files": [],
        "indexed_at": utc_now_iso(),
    }

    for path in files:
        ap = str(path.resolve())
        seen_paths.add(ap)
        stat = path.stat()
        rec = existing.get(ap)
        try:
            if rec and int(rec["mtime_ns"]) == stat.st_mtime_ns and int(rec["size_bytes"]) == stat.st_size:
                report["skipped"] += 1
                continue

            sha = compute_sha256(path)
            if rec and rec.get("sha256") == sha:
                conn.execute(
                    """
                    UPDATE kb_files
                    SET mtime_ns=?, size_bytes=?, indexed_at=?
                    WHERE path=?
                    """,
                    (stat.st_mtime_ns, stat.st_size, utc_now_iso(), ap),
                )
                report["metadata_only"] += 1
                report["metadata_only_paths"].append(ap)
                continue

            parser, text = extract_text(path)
            chunks = _chunk_text(text)
            source_group = infer_source_group(path, kb_root)

            conn.execute("DELETE FROM kb_chunks WHERE path=?", (ap,))
            for idx, chunk in enumerate(chunks):
                conn.execute(
                    "INSERT INTO kb_chunks(path, source_group, chunk_index, text) VALUES(?,?,?,?)",
                    (ap, source_group, idx, chunk),
                )

            conn.execute(
                """
                INSERT INTO kb_files(path, mtime_ns, size_bytes, sha256, parser, source_group, chunk_count, indexed_at)
                VALUES(?,?,?,?,?,?,?,?)
                ON CONFLICT(path) DO UPDATE SET
                    mtime_ns=excluded.mtime_ns,
                    size_bytes=excluded.size_bytes,
                    sha256=excluded.sha256,
                    parser=excluded.parser,
                    source_group=excluded.source_group,
                    chunk_count=excluded.chunk_count,
                    indexed_at=excluded.indexed_at
                """,
                (ap, stat.st_mtime_ns, stat.st_size, sha, parser, source_group, len(chunks), utc_now_iso()),
            )

            report["files"].append(
                {
                    "path": ap,
                    "parser": parser,
                    "source_group": source_group,
                    "chunk_count": len(chunks),
                }
            )
            if rec:
                report["updated"] += 1
            else:
                report["created"] += 1
        except Exception as exc:  # pragma: no cover - keeps sync resilient
            report["errors"].append({"path": ap, "error": str(exc)})

    for old_path in existing:
        if old_path in seen_paths:
            continue
        conn.execute("DELETE FROM kb_chunks WHERE path=?", (old_path,))
        conn.execute("DELETE FROM kb_files WHERE path=?", (old_path,))
        report["removed"] += 1

    conn.commit()
    report["ok"] = len(report["errors"]) == 0

    if report_path:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json_dumps(report), encoding="utf-8")

    return report


def sync_kb_with_rag(
    *,
    conn: sqlite3.Connection,
    kb_root: Path,
    report_path: Path | None = None,
    rag_backend: str = "clawrag",
    rag_base_url: str = "http://127.0.0.1:8080",
    rag_collection: str = "translation-kb",
) -> dict[str, Any]:
    local_report = sync_kb(conn=conn, kb_root=kb_root, report_path=report_path)
    rag_report: dict[str, Any] = {"ok": False, "backend": "local", "mode": "disabled"}
    if str(rag_backend).strip().lower() == "clawrag":
        changed_paths = [str(x.get("path")) for x in (local_report.get("files") or []) if str(x.get("path", "")).strip()]
        changed_paths.extend([str(p) for p in (local_report.get("metadata_only_paths") or []) if str(p).strip()])
        changed_paths = sorted(set([str(p).strip() for p in changed_paths if str(p).strip()]))
        rag_report = clawrag_sync(
            changed_paths=changed_paths,
            base_url=rag_base_url,
            collection=rag_collection,
        )
    return {"ok": local_report.get("ok", False), "local_report": local_report, "rag_report": rag_report}


def retrieve_kb(
    *,
    conn: sqlite3.Connection,
    query: str,
    task_type: str = "",
    top_k: int = 8,
    kb_root: Path | None = None,
    kb_company: str = "",
    isolation_mode: str = "reference_only",
) -> list[dict[str, Any]]:
    tokens = [t.lower() for t in re.findall(r"[A-Za-z0-9\u0600-\u06FF_]+", query) if len(t) >= 2]
    if not tokens:
        return []

    task = task_type.upper().strip()
    task_boosts = {
        "REVISION_UPDATE": {"glossary": 1.25, "previously_translated": 1.2},
        "NEW_TRANSLATION": {"glossary": 1.2, "arabic_source": 1.1},
        "BILINGUAL_REVIEW": {"glossary": 1.2, "translated_en": 1.15},
        "EN_ONLY_EDIT": {"translated_en": 1.2, "previously_translated": 1.1},
        "MULTI_FILE_BATCH": {"glossary": 1.15, "previously_translated": 1.1},
        "TERMINOLOGY_ENFORCEMENT": {"glossary": 1.4, "translated_en": 1.15},
        "FORMAT_CRITICAL_TASK": {"glossary": 1.2, "previously_translated": 1.15},
        "LOW_CONTEXT_TASK": {"glossary": 1.1},
    }.get(task, {})

    if _ensure_kb_fts(conn):
        match_query = " OR ".join(sorted(set(tokens)))
        where_sql = ""
        where_params: list[Any] = []
        if kb_root and kb_company.strip():
            ref_like, company_like = _reference_like_filters(kb_root=kb_root, kb_company=kb_company.strip())
            mode = (isolation_mode or "reference_only").strip().lower()
            if mode == "all":
                where_sql = " AND c.path LIKE ?"
                where_params.append(company_like)
            else:
                where_sql = " AND (c.path NOT LIKE ? OR c.path LIKE ?)"
                where_params.extend([ref_like, company_like])

        rows = conn.execute(
            f"""
            SELECT
              c.path AS path,
              c.source_group AS source_group,
              c.chunk_index AS chunk_index,
              substr(c.text, 1, 700) AS snippet,
              bm25(kb_chunks_fts) AS rank
            FROM kb_chunks_fts
            JOIN kb_chunks c ON c.id = kb_chunks_fts.rowid
            WHERE kb_chunks_fts MATCH ?{where_sql}
            LIMIT ?
            """,
            [match_query, *where_params, max(10, int(top_k) * 8)],
        ).fetchall()

        scored: list[dict[str, Any]] = []
        for row in rows:
            source_group = row["source_group"] or "general"
            base = SOURCE_GROUP_WEIGHTS.get(source_group, SOURCE_GROUP_WEIGHTS["general"])
            boost = task_boosts.get(source_group, 1.0)
            raw_rank = float(row["rank"] or 0.0)
            inv = 1.0 / (1.0 + max(0.0, raw_rank))
            score = round(inv * float(base) * float(boost), 6)
            scored.append(
                {
                    "path": row["path"],
                    "source_group": source_group,
                    "chunk_index": int(row["chunk_index"]),
                    "snippet": str(row["snippet"] or ""),
                    "score": score,
                }
            )
        scored.sort(key=lambda x: x["score"], reverse=True)
        return scored[: max(1, int(top_k))]

    where_sql = ""
    where_params: list[Any] = []
    if kb_root and kb_company.strip():
        ref_like, company_like = _reference_like_filters(kb_root=kb_root, kb_company=kb_company.strip())
        mode = (isolation_mode or "reference_only").strip().lower()
        if mode == "all":
            where_sql = " WHERE path LIKE ?"
            where_params.append(company_like)
        else:
            where_sql = " WHERE (path NOT LIKE ? OR path LIKE ?)"
            where_params.extend([ref_like, company_like])

    rows = conn.execute(f"SELECT path, source_group, chunk_index, text FROM kb_chunks{where_sql}", tuple(where_params)).fetchall()
    scored: list[dict[str, Any]] = []
    for row in rows:
        text = (row["text"] or "").lower()
        if not text:
            continue
        match_hits = 0
        for tk in tokens:
            match_hits += text.count(tk)
        if match_hits <= 0:
            continue
        source_group = row["source_group"] or "general"
        base = SOURCE_GROUP_WEIGHTS.get(source_group, SOURCE_GROUP_WEIGHTS["general"])
        boost = task_boosts.get(source_group, 1.0)
        score = round(float(match_hits) * base * boost, 4)
        scored.append(
            {
                "path": row["path"],
                "source_group": source_group,
                "chunk_index": int(row["chunk_index"]),
                "snippet": row["text"][:700],
                "score": score,
            }
        )

    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[: max(1, int(top_k))]


def retrieve_kb_with_fallback(
    *,
    conn: sqlite3.Connection,
    query: str,
    task_type: str = "",
    kb_root: Path | None = None,
    kb_company: str = "",
    isolation_mode: str = "reference_only",
    rag_backend: str = "clawrag",
    rag_base_url: str = "http://127.0.0.1:8080",
    rag_collection: str = "translation-kb",
    top_k_clawrag: int = 12,
    top_k_local: int = 8,
) -> dict[str, Any]:
    q = (query or "").strip()
    if not q:
        return {"backend": "local", "hits": [], "status_flags": [], "rag_result": {"ok": True, "detail": "empty_query"}}

    if str(rag_backend).strip().lower() == "clawrag":
        rag_result = clawrag_search(
            query=q,
            top_k=max(1, int(top_k_clawrag)),
            base_url=rag_base_url,
            collection=rag_collection,
        )
        rag_hits = list(rag_result.get("hits") or [])
        if rag_result.get("ok") and rag_hits:
            filtered = [
                h for h in rag_hits
                if _allow_kb_path(
                    str(h.get("path") or ""),
                    kb_root=kb_root,
                    kb_company=kb_company,
                    isolation_mode=isolation_mode,
                )
            ]
            if filtered:
                return {
                    "backend": "clawrag",
                    "hits": filtered[: max(1, int(top_k_clawrag))],
                    "status_flags": [],
                    "rag_result": rag_result,
                }
        local_hits = retrieve_kb(
            conn=conn,
            query=q,
            task_type=task_type,
            top_k=max(1, int(top_k_local)),
            kb_root=kb_root,
            kb_company=kb_company,
            isolation_mode=isolation_mode,
        )
        return {
            "backend": "local",
            "hits": local_hits,
            "status_flags": ["rag_fallback_local"] + (["rag_filtered_empty"] if rag_result.get("ok") and rag_hits else []),
            "rag_result": rag_result,
        }

    local_hits = retrieve_kb(
        conn=conn,
        query=q,
        task_type=task_type,
        top_k=max(1, int(top_k_local)),
        kb_root=kb_root,
        kb_company=kb_company,
        isolation_mode=isolation_mode,
    )
    return {"backend": "local", "hits": local_hits, "status_flags": [], "rag_result": {"ok": True, "mode": "local_only"}}
