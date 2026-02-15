#!/usr/bin/env python3
"""DOCX format-preserving translation support.

Goal: keep original Word formatting/structure while replacing only the text content.

This module keeps the original structure and uses a best-effort approach to preserve
run-level formatting by reusing existing runs and distributing replacement text across them.
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\u00A0", " ")).strip()


@dataclass(frozen=True)
class DocxUnit:
    unit_id: str
    kind: str  # "paragraph" | "table_cell"
    style: str
    text: str


def extract_units(
    docx_path: Path,
    *,
    include_tables: bool = True,
    max_units: int | None = None,
    max_chars_per_unit: int = 800,
) -> tuple[list[DocxUnit], dict[str, Any]]:
    """Extract translatable units from a DOCX as stable IDs.

    IDs are based on body child index + table position:
    - Paragraph: p:<body_index>
    - Table cell: t<table_index>:r<row>:c<col> (1-based)
    """
    docx_path = Path(docx_path).expanduser().resolve()
    doc = Document(str(docx_path))
    units: list[DocxUnit] = []
    truncated = False

    block_index = 0
    table_index = 0
    for child in doc.element.body.iterchildren():
        block_index += 1
        if isinstance(child, CT_P):
            para = Paragraph(child, doc)
            text = _normalize_text(para.text)
            if not text:
                continue
            if max_chars_per_unit > 0 and len(text) > max_chars_per_unit:
                text = text[:max_chars_per_unit]
            units.append(
                DocxUnit(
                    unit_id=f"p:{block_index}",
                    kind="paragraph",
                    style=para.style.name if para.style else "",
                    text=text,
                )
            )
        elif include_tables and isinstance(child, CT_Tbl):
            table_index += 1
            table = Table(child, doc)
            for r_idx, row in enumerate(table.rows, start=1):
                for c_idx, cell in enumerate(row.cells, start=1):
                    text = _normalize_text(cell.text.replace("\n", " / "))
                    if not text:
                        continue
                    if max_chars_per_unit > 0 and len(text) > max_chars_per_unit:
                        text = text[:max_chars_per_unit]
                    units.append(
                        DocxUnit(
                            unit_id=f"t{table_index}:r{r_idx}:c{c_idx}",
                            kind="table_cell",
                            style="",
                            text=text,
                        )
                    )
        if max_units is not None and len(units) >= max_units:
            truncated = True
            break

    return units, {
        "file": docx_path.name,
        "unit_count": len(units),
        "truncated": truncated,
        "max_units": max_units,
    }


def units_to_payload(units: Iterable[DocxUnit]) -> list[dict[str, Any]]:
    return [
        {
            "id": u.unit_id,
            "kind": u.kind,
            "style": u.style,
            "text": u.text,
        }
        for u in units
    ]


def _normalize_docx_translation_map(entries: Any) -> dict[str, str]:
    """Normalize model output into {unit_id: text}."""
    out: dict[str, str] = {}
    if not entries:
        return out
    if isinstance(entries, dict):
        for k, v in entries.items():
            key = str(k or "").strip()
            if not key:
                continue
            out[key] = str(v or "")
        return out
    if not isinstance(entries, list):
        return out
    for item in entries:
        if not isinstance(item, dict):
            continue
        unit_id = str(item.get("id") or item.get("unit_id") or item.get("block_id") or item.get("cell_id") or "").strip()
        if not unit_id:
            continue
        out[unit_id] = str(item.get("text") or "")
    return out


def apply_translation_map(
    *,
    template_docx: Path,
    output_docx: Path,
    translation_map_entries: Any,
) -> dict[str, Any]:
    """Copy template docx and replace unit text by IDs."""
    template_docx = Path(template_docx).expanduser().resolve()
    output_docx = Path(output_docx).expanduser().resolve()
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(str(template_docx), str(output_docx))

    mapped = _normalize_docx_translation_map(translation_map_entries)

    def _split_by_runs(*, runs: list[Any], new_text: str) -> list[str]:
        """Split new_text into len(runs) segments proportional to original run text lengths."""
        count = len(runs)
        if count <= 0:
            return []
        if not new_text:
            return [""] * count

        lengths: list[int] = []
        total = 0
        for run in runs:
            try:
                ln = len(str(getattr(run, "text", "") or ""))
            except Exception:
                ln = 0
            lengths.append(max(0, int(ln)))
            total += lengths[-1]

        if total <= 0:
            segs = [""] * count
            segs[0] = str(new_text)
            return segs

        new_len = len(new_text)
        bases: list[int] = []
        remainders: list[tuple[float, int]] = []
        base_sum = 0
        for idx, ln in enumerate(lengths):
            raw = (new_len * ln) / total
            base = int(raw)
            bases.append(base)
            base_sum += base
            remainders.append((raw - base, idx))

        remaining = max(0, new_len - base_sum)
        remainders.sort(reverse=True)
        for i in range(min(remaining, count)):
            bases[remainders[i][1]] += 1

        segments: list[str] = []
        cursor = 0
        for ln in bases:
            segments.append(new_text[cursor : cursor + ln])
            cursor += ln
        if cursor < new_len:
            segments[-1] += new_text[cursor:]
        return segments

    def _replace_paragraph_text_preserving_runs(paragraph: Paragraph, new_text: str) -> None:
        runs = list(paragraph.runs or [])
        if not runs:
            paragraph.add_run(str(new_text))
            return
        segments = _split_by_runs(runs=runs, new_text=str(new_text))
        for run, seg in zip(runs, segments):
            try:
                run.text = seg
            except Exception:  # pragma: no cover
                pass

    def _replace_table_cell_text_preserving_paragraphs(cell: Any, new_text: str) -> None:
        paragraphs = list(getattr(cell, "paragraphs", []) or [])
        if not paragraphs:
            # Fallback: avoid structural changes when possible, but cell.text is better than losing content.
            try:
                cell.text = str(new_text)
            except Exception:  # pragma: no cover
                pass
            return

        lines = [ln.rstrip("\r") for ln in str(new_text).split("\n")]
        if not lines:
            lines = [""]
        if len(lines) > len(paragraphs):
            extra = " ".join(lines[len(paragraphs):]).strip()
            lines = lines[: len(paragraphs)]
            if extra:
                lines[-1] = (lines[-1] + " " + extra).strip()

        for idx, para in enumerate(paragraphs):
            value = lines[idx] if idx < len(lines) else ""
            _replace_paragraph_text_preserving_runs(para, value)

    doc = Document(str(output_docx))
    applied = 0
    block_index = 0
    table_index = 0
    for child in doc.element.body.iterchildren():
        block_index += 1
        if isinstance(child, CT_P):
            unit_id = f"p:{block_index}"
            if unit_id not in mapped:
                continue
            para = Paragraph(child, doc)
            _replace_paragraph_text_preserving_runs(para, str(mapped[unit_id]))
            applied += 1
        elif isinstance(child, CT_Tbl):
            table_index += 1
            table = Table(child, doc)
            for r_idx, row in enumerate(table.rows, start=1):
                for c_idx, cell in enumerate(row.cells, start=1):
                    unit_id = f"t{table_index}:r{r_idx}:c{c_idx}"
                    if unit_id not in mapped:
                        continue
                    _replace_table_cell_text_preserving_paragraphs(cell, str(mapped[unit_id]))
                    applied += 1

    doc.save(str(output_docx))
    return {
        "ok": True,
        "template": str(template_docx),
        "output": str(output_docx),
        "applied_count": applied,
    }
