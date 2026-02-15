#!/usr/bin/env python3
"""Write V5.2 artifact bundle into the verify folder."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook

from scripts.compose_docx_from_draft import build_doc
from scripts.docx_preserver import apply_translation_map as apply_docx_translation_map
from scripts.xlsx_preserver import apply_translation_map as apply_xlsx_translation_map

SYSTEM_DIR_NAME = ".system"


def _write_text(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value, encoding="utf-8")


def _write_json(path: Path, value: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def _text_to_lines(text: str) -> list[str]:
    return [line.rstrip() for line in (text or "").splitlines()]


def _write_docx(path: Path, title: str, lines: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=1)
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue
        doc.add_paragraph(stripped)
    doc.save(str(path))


def _write_xlsx(path: Path, *, final_text: str, change_log_points: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Final"
    for idx, line in enumerate([ln for ln in final_text.splitlines() if ln.strip()], start=1):
        ws.cell(row=idx, column=1, value=line.strip())

    log = wb.create_sheet(title="ChangeLog")
    if change_log_points:
        for idx, row in enumerate(change_log_points, start=1):
            log.cell(row=idx, column=1, value=str(row))
    else:
        log.cell(row=1, column=1, value="No explicit change log points were returned by model.")
    wb.save(str(path))


def build_review_brief_lines(
    *,
    task_type: str,
    quality_report: dict[str, Any],
    status_flags: list[str],
    review_questions: list[str],
) -> list[str]:
    rounds = quality_report.get("rounds", [])
    convergence = bool(quality_report.get("convergence_reached"))
    stop_reason = quality_report.get("stop_reason", "unknown")
    lines = [
        f"Task type: {task_type}",
        f"Convergence reached: {convergence}",
        f"Stop reason: {stop_reason}",
        f"Status flags: {', '.join(status_flags) if status_flags else 'none'}",
        "",
        "## Round Results",
    ]
    if not rounds:
        lines.append("- No rounds produced.")
    else:
        for item in rounds:
            lines.append(
                f"- Round {item.get('round')}: pass={item.get('pass')} codex={item.get('codex_pass')} gemini={item.get('gemini_pass')}"
            )
            unresolved = item.get("unresolved") or []
            if unresolved:
                lines.append(f"-   unresolved: {', '.join(unresolved)}")

    lines.extend(["", "## Questions / Notes"])
    if review_questions:
        for q in review_questions:
            lines.append(f"- {q}")
    else:
        lines.append("- No extra notes.")
    lines.extend(
        [
            "",
            "## Manual Policy",
            "- Output is in _VERIFY only.",
            "- System will not auto-move files to final folder.",
            "- After manual validation, use command: ok (status only).",
        ]
    )
    return lines


def _ensure_change_log_text(change_log_points: list[str], task_type: str) -> str:
    if change_log_points:
        body = "\n".join(f"- {x}" for x in change_log_points if str(x).strip())
    else:
        body = "- No explicit change log points were returned by model."
    return "\n".join(
        [
            f"# Change Log ({task_type})",
            "",
            body,
            "",
            "## Delivery Note",
            "- This artifact is generated in _VERIFY only.",
            "- Manual move is required for final delivery.",
        ]
    )


def write_artifacts(
    *,
    review_dir: str,
    draft_a_template_path: str | None,
    delta_pack: dict[str, Any],
    model_scores: dict[str, Any],
    quality: dict[str, Any],
    quality_report: dict[str, Any],
    job_id: str,
    task_type: str,
    confidence: float,
    estimated_minutes: int,
    runtime_timeout_minutes: int,
    iteration_count: int,
    double_pass: bool,
    status_flags: list[str],
    candidate_files: list[dict[str, Any]],
    review_questions: list[str],
    draft_payload: dict[str, Any] | None = None,
    generate_final_xlsx: bool = False,
    plan_payload: dict[str, Any] | None = None,
) -> dict[str, Any]:
    review = Path(review_dir)
    system = review / SYSTEM_DIR_NAME
    review.mkdir(parents=True, exist_ok=True)
    system.mkdir(parents=True, exist_ok=True)

    draft_payload = draft_payload or {}
    plan_payload = plan_payload or {}

    final_text = str(
        draft_payload.get("final_text")
        or draft_payload.get("draft_a_text")
        or draft_payload.get("draft_b_text")
        or ""
    )
    final_reflow_text = str(
        draft_payload.get("final_reflow_text")
        or draft_payload.get("draft_b_text")
        or final_text
        or ""
    )
    review_points = [str(x) for x in (draft_payload.get("review_brief_points") or [])]
    change_log_points = [str(x) for x in (draft_payload.get("change_log_points") or [])]
    docx_translation_map = draft_payload.get("docx_translation_map") or []
    xlsx_translation_map = draft_payload.get("xlsx_translation_map") or []

    final_docx = review / "Final.docx"
    final_reflow_docx = review / "Final-Reflow.docx"
    review_brief_docx = review / "Review Brief.docx"
    change_log_md = review / "Change Log.md"
    final_xlsx = review / "Final.xlsx"

    execution_plan_json = system / "execution_plan.json"
    quality_report_json = system / "quality_report.json"
    delta_summary_json = system / "Delta Summary.json"
    model_scores_json = system / "Model Scores.json"

    template = Path(draft_a_template_path) if draft_a_template_path else None
    if template and template.exists() and docx_translation_map:
        apply_docx_translation_map(template_docx=template, output_docx=final_docx, translation_map_entries=docx_translation_map)
    elif template and template.exists():
        build_doc(template, final_docx, final_text)
    else:
        _write_docx(final_docx, "Final", _text_to_lines(final_text))

    _write_docx(final_reflow_docx, "Final-Reflow", _text_to_lines(final_reflow_text))

    review_lines = build_review_brief_lines(
        task_type=task_type,
        quality_report=quality_report,
        status_flags=status_flags,
        review_questions=(review_points + review_questions),
    )
    _write_docx(review_brief_docx, "Review Brief", review_lines)

    _write_text(change_log_md, _ensure_change_log_text(change_log_points, task_type))

    xlsx_sources = [
        Path(str(item.get("path") or "")).expanduser().resolve()
        for item in (candidate_files or [])
        if str(item.get("path") or "") and Path(str(item.get("path") or "")).suffix.lower() == ".xlsx"
    ]

    beautify_xlsx = str((plan_payload.get("meta") or {}).get("beautify_xlsx", "")).strip()
    if not beautify_xlsx:
        beautify_xlsx = "1"
    beautify_xlsx_enabled = beautify_xlsx not in {"0", "false", "off", "no"}

    xlsx_entries: list[dict[str, Any]] = []
    if generate_final_xlsx and xlsx_sources and xlsx_translation_map:
        if len(xlsx_sources) == 1:
            src = xlsx_sources[0]
            res = apply_xlsx_translation_map(
                source_xlsx=src,
                output_xlsx=final_xlsx,
                translation_map_entries=xlsx_translation_map,
                beautify=beautify_xlsx_enabled,
            )
            xlsx_entries.append({"name": final_xlsx.name, "path": str(final_xlsx.resolve()), "source_path": str(src), "apply_result": res})
        else:
            for src in xlsx_sources:
                out_path = review / f"{src.stem}_translated.xlsx"
                res = apply_xlsx_translation_map(
                    source_xlsx=src,
                    output_xlsx=out_path,
                    translation_map_entries=xlsx_translation_map,
                    beautify=beautify_xlsx_enabled,
                )
                xlsx_entries.append({"name": out_path.name, "path": str(out_path.resolve()), "source_path": str(src), "apply_result": res})
    elif generate_final_xlsx:
        _write_xlsx(final_xlsx, final_text=final_text, change_log_points=change_log_points)

    plan_write = {
        "job_id": job_id,
        "task_type": task_type,
        "confidence": confidence,
        "estimated_minutes": estimated_minutes,
        "runtime_timeout_minutes": runtime_timeout_minutes,
        "iteration_count": iteration_count,
        "double_pass": double_pass,
        "status_flags": status_flags,
        "candidate_files": candidate_files,
        "pipeline_version": str(
            (plan_payload.get("plan") or {}).get("pipeline_version")
            or (plan_payload.get("meta") or {}).get("pipeline_version")
            or ""
        ).strip(),
        "markdown_policy": (plan_payload.get("meta") or {}).get("markdown_policy") or {},
        "vision_policy": (plan_payload.get("meta") or {}).get("vision_policy") or {},
        "plan_payload": plan_payload,
    }
    _write_json(execution_plan_json, plan_write)
    _write_json(quality_report_json, quality_report)
    _write_json(delta_summary_json, delta_pack)
    _write_json(model_scores_json, model_scores)

    manifest: dict[str, Any] = {
        "final_docx": str(final_docx.resolve()),
        "final_reflow_docx": str(final_reflow_docx.resolve()),
        "review_brief_docx": str(review_brief_docx.resolve()),
        "change_log_md": str(change_log_md.resolve()),
        "execution_plan_json": str(execution_plan_json.resolve()),
        "quality_report_json": str(quality_report_json.resolve()),
        "delta_summary_json": str(delta_summary_json.resolve()),
        "model_scores_json": str(model_scores_json.resolve()),
    }
    if xlsx_entries and len(xlsx_entries) == 1 and xlsx_entries[0]["name"] == final_xlsx.name:
        manifest["final_xlsx"] = str(final_xlsx.resolve())
        manifest["xlsx_files"] = xlsx_entries
    elif xlsx_entries:
        manifest["xlsx_files"] = xlsx_entries
    elif generate_final_xlsx:
        manifest["final_xlsx"] = str(final_xlsx.resolve())
    return manifest
