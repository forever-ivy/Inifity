#!/usr/bin/env python3

import tempfile
import unittest
from pathlib import Path

from docx import Document

from scripts.docx_preserver import apply_translation_map


def _make_docx(path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.add_run("Hello ")
    run2 = p.add_run("World")
    run2.bold = True

    table = doc.add_table(rows=1, cols=1)
    cell_p = table.cell(0, 0).paragraphs[0]
    cell_run = cell_p.add_run("Cell")
    cell_run.italic = True

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


class DocxPreserverTest(unittest.TestCase):
    def test_apply_translation_preserves_run_formatting_best_effort(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            out = Path(tmp) / "out.docx"
            _make_docx(template)

            res = apply_translation_map(
                template_docx=template,
                output_docx=out,
                translation_map_entries=[
                    {"id": "p:1", "text": "Bonjour Monde"},
                    {"id": "t1:r1:c1", "text": "Cellule"},
                ],
            )
            self.assertTrue(res.get("ok"))
            self.assertEqual(res.get("applied_count"), 2)

            doc = Document(str(out))
            self.assertGreaterEqual(len(doc.paragraphs), 1)
            para = doc.paragraphs[0]
            self.assertEqual(para.style.name, "Heading 1")
            self.assertEqual(para.text, "Bonjour Monde")
            self.assertEqual(len(para.runs), 2)
            self.assertTrue(para.runs[1].bold)

            self.assertEqual(len(doc.tables), 1)
            cell_para = doc.tables[0].cell(0, 0).paragraphs[0]
            self.assertEqual(cell_para.text, "Cellule")
            self.assertTrue(cell_para.runs[0].italic)


if __name__ == "__main__":
    unittest.main()

