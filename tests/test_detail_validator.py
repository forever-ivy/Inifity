#!/usr/bin/env python3
"""Tests for detail_validator.py module."""

import json
import tempfile
import unittest
from pathlib import Path

from scripts.detail_validator import (
    Category,
    DocxStructureValidator,
    Severity,
    ValidationConfig,
    ValidationIssue,
    ValidationResult,
    ValidationReportGenerator,
    ValidationReportGenerator,
    XlsxStructureValidator,
    validate_file_pair,
    validate_job_artifacts,
)

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False
    Pt = None  # type: ignore

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    OPENPYXL_AVAILABLE = True
except Exception:
    OPENPYXL_AVAILABLE = False


def _make_test_docx(path: Path, with_table: bool = False) -> None:
    """Create a test DOCX file."""
    doc = Document()

    # Add heading with specific font
    p = doc.add_heading("Test Document", level=1)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(16)

    # Add normal paragraph
    p = doc.add_paragraph("This is a test paragraph with some text.")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    if with_table:
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _make_test_xlsx(path: Path, with_merged: bool = False) -> None:
    """Create a test XLSX file."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TestSheet"

    # Add cells with formatting
    ws["A1"] = "Header"
    ws["A1"].font = Font(bold=True, size=12, color="FF0000")
    ws["A1"].fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

    ws["A2"] = "Data"
    ws["A2"].font = Font(size=10)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center")

    ws.column_dimensions["A"].width = 20
    ws.row_dimensions[1].height = 25

    if with_merged:
        ws.merge_cells("B1:C1")
        ws["B1"] = "Merged Header"

    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))
    wb.close()


class ValidationIssueTest(unittest.TestCase):
    def test_to_dict(self) -> None:
        issue = ValidationIssue(
            category=Category.FONT,
            severity=Severity.WARNING,
            location="p:1",
            element_type="paragraph",
            expected="font: Arial 11pt",
            actual="font: Calibri 10pt",
            hint="Change font to Arial 11pt",
        )
        result = issue.to_dict()
        self.assertEqual(result["category"], "font")
        self.assertEqual(result["severity"], "warning")
        self.assertEqual(result["location"], "p:1")
        self.assertIn("fix_hint", result)


class ValidationResultTest(unittest.TestCase):
    def test_add_issue(self) -> None:
        result = ValidationResult(
            file_name="test.docx",
            file_path="/tmp/test.docx",
            format_type="docx",
            valid=True,
        )

        result.add_issue(ValidationIssue(
            category=Category.FONT,
            severity=Severity.CRITICAL,
            location="p:1",
            element_type="paragraph",
            expected="Arial",
            actual="Calibri",
            hint="Fix font",
        ))

        self.assertEqual(result.total_checks, 1)
        self.assertEqual(result.failed, 1)
        self.assertEqual(len(result.issues), 1)

    def test_calculate_score(self) -> None:
        result = ValidationResult(
            file_name="test.docx",
            file_path="/tmp/test.docx",
            format_type="docx",
            valid=True,
        )

        # Add 2 critical, 3 warnings
        for _ in range(2):
            result.add_issue(ValidationIssue(
                category=Category.FONT,
                severity=Severity.CRITICAL,
                location="p:1",
                element_type="paragraph",
                expected="",
                actual="",
                hint="",
            ))
        for _ in range(3):
            result.add_issue(ValidationIssue(
                category=Category.STYLE,
                severity=Severity.WARNING,
                location="p:1",
                element_type="paragraph",
                expected="",
                actual="",
                hint="",
            ))

        result.calculate_score(critical_weight=0.15, warning_weight=0.05)
        # Score = 1 - (2 * 0.15 + 3 * 0.05) = 1 - 0.45 = 0.55
        self.assertAlmostEqual(result.format_fidelity_score, 0.55)
        self.assertFalse(result.valid)  # Not valid due to critical issues

    def test_to_dict(self) -> None:
        result = ValidationResult(
            file_name="test.docx",
            file_path="/tmp/test.docx",
            format_type="docx",
            valid=True,
        )
        result.add_pass()
        result.calculate_score()

        d = result.to_dict()
        self.assertEqual(d["file_name"], "test.docx")
        self.assertEqual(d["format_type"], "docx")
        self.assertIn("score", d)


@unittest.skipIf(not DOCX_AVAILABLE, "python-docx not available")
class DocxValidatorTest(unittest.TestCase):
    def test_validate_identical_documents(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            original = tmp_path / "original.docx"
            translated = tmp_path / "translated.docx"

            _make_test_docx(original)
            _make_test_docx(translated)

            validator = DocxStructureValidator(ValidationConfig())
            result = validator.validate(original, translated)

            self.assertEqual(result.format_type, "docx")
            self.assertGreaterEqual(result.format_fidelity_score, 0.9)

    def test_detects_font_difference(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            original = tmp_path / "original.docx"
            translated = tmp_path / "translated.docx"

            _make_test_docx(original)

            # Create translated with different font
            doc = Document()
            p = doc.add_heading("Test Document", level=1)
            for run in p.runs:
                run.font.name = "Times New Roman"  # Different font
                run.font.size = Pt(16)

            p = doc.add_paragraph("This is a test paragraph with some text.")
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

            doc.save(str(translated))

            validator = DocxStructureValidator(ValidationConfig(
                docx_check_fonts=True,
                docx_font_size_delta=0.1,
            ))
            result = validator.validate(original, translated)

            self.assertGreater(len(result.issues), 0)
            font_issues = [i for i in result.issues if i.category == Category.FONT]
            self.assertGreater(len(font_issues), 0)

    def test_detects_table_dimension_difference(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            original = tmp_path / "original.docx"
            translated = tmp_path / "translated.docx"

            _make_test_docx(original, with_table=True)

            # Create translated with different table
            doc = Document()
            table = doc.add_table(rows=3, cols=2)  # Different row count
            for i in range(3):
                for j in range(2):
                    table.cell(i, j).text = f"Cell {i},{j}"
            doc.save(str(translated))

            validator = DocxStructureValidator(ValidationConfig(
                docx_check_tables=True,
            ))
            result = validator.validate(original, translated)

            table_issues = [i for i in result.issues if i.category == Category.TABLE]
            self.assertGreater(len(table_issues), 0)


@unittest.skipIf(not OPENPYXL_AVAILABLE, "openpyxl not available")
class XlsxValidatorTest(unittest.TestCase):
    def test_validate_identical_workbooks(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            original = tmp_path / "original.xlsx"
            translated = tmp_path / "translated.xlsx"

            _make_test_xlsx(original)
            _make_test_xlsx(translated)

            validator = XlsxStructureValidator(ValidationConfig())
            result = validator.validate(original, translated)

            self.assertEqual(result.format_type, "xlsx")
            self.assertGreaterEqual(result.format_fidelity_score, 0.9)

    def test_detects_font_difference(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            original = tmp_path / "original.xlsx"
            translated = tmp_path / "translated.xlsx"

            _make_test_xlsx(original)

            # Create translated with different font
            wb = openpyxl.Workbook()
            ws = wb.active
            ws["A1"] = "Header"
            ws["A1"].font = Font(bold=False, size=10)  # Different
            wb.save(str(translated))
            wb.close()

            validator = XlsxStructureValidator(ValidationConfig(
                xlsx_check_fonts=True,
            ))
            result = validator.validate(original, translated)

            font_issues = [i for i in result.issues if i.category == Category.CELL_FONT]
            self.assertGreater(len(font_issues), 0)

    def test_detects_merged_region_difference(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            original = tmp_path / "original.xlsx"
            translated = tmp_path / "translated.xlsx"

            _make_test_xlsx(original, with_merged=True)

            # Create translated WITHOUT merge (use same sheet name)
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "TestSheet"  # Match the original sheet name
            ws["A1"] = "Header"
            ws["A1"].font = Font(bold=True, size=12)
            ws["A1"].fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
            ws["A2"] = "Data"
            ws["A2"].font = Font(size=10)
            ws["A2"].alignment = Alignment(horizontal="left", vertical="center")
            ws.column_dimensions["A"].width = 20
            ws.row_dimensions[1].height = 25
            # No merge in translated
            wb.save(str(translated))
            wb.close()

            validator = XlsxStructureValidator(ValidationConfig(
                xlsx_check_merged=True,
            ))
            result = validator.validate(original, translated)

            merge_issues = [i for i in result.issues if i.category == Category.MERGED_REGIONS]
            self.assertGreater(len(merge_issues), 0)


class ValidationReportGeneratorTest(unittest.TestCase):
    def test_generate_markdown(self) -> None:
        result = ValidationResult(
            file_name="test.docx",
            file_path="/tmp/test.docx",
            format_type="docx",
            valid=True,
        )
        result.add_issue(ValidationIssue(
            category=Category.FONT,
            severity=Severity.CRITICAL,
            location="p:1",
            element_type="paragraph",
            expected="Arial 11pt",
            actual="Calibri 10pt",
            hint="Change font to Arial 11pt",
        ))
        result.calculate_score()

        generator = ValidationReportGenerator()
        report = generator.generate_markdown([result])

        self.assertIn("# Detail Validation Report", report)
        self.assertIn("test.docx", report)
        self.assertIn("Arial 11pt", report)
        self.assertIn("Calibri 10pt", report)

    def test_generate_summary(self) -> None:
        result1 = ValidationResult(
            file_name="test1.docx",
            file_path="/tmp/test1.docx",
            format_type="docx",
            valid=True,
        )
        result1.add_pass()
        result1.calculate_score()

        result2 = ValidationResult(
            file_name="test2.docx",
            file_path="/tmp/test2.docx",
            format_type="docx",
            valid=False,
        )
        result2.add_issue(ValidationIssue(
            category=Category.FONT,
            severity=Severity.CRITICAL,
            location="p:1",
            element_type="paragraph",
            expected="",
            actual="",
            hint="",
        ))
        result2.calculate_score()

        generator = ValidationReportGenerator()
        summary = generator.generate_summary([result1, result2])

        self.assertEqual(summary["total_files"], 2)
        self.assertEqual(summary["files_valid"], 1)
        self.assertIn("score", summary)

    def test_extract_fix_hints(self) -> None:
        result = ValidationResult(
            file_name="test.docx",
            file_path="/tmp/test.docx",
            format_type="docx",
            valid=False,
        )
        result.add_issue(ValidationIssue(
            category=Category.FONT,
            severity=Severity.CRITICAL,
            location="p:1",
            element_type="paragraph",
            expected="",
            actual="",
            hint="Fix the font in paragraph 1",
        ))
        result.add_issue(ValidationIssue(
            category=Category.STYLE,
            severity=Severity.WARNING,
            location="p:2",
            element_type="paragraph",
            expected="",
            actual="",
            hint="Apply heading style to paragraph 2",
        ))

        generator = ValidationReportGenerator()
        hints = generator.extract_fix_hints([result])

        self.assertEqual(len(hints), 2)
        self.assertTrue(any("Fix the font" in h for h in hints))
        self.assertTrue(any("Apply heading style" in h for h in hints))


class ValidateFilePairTest(unittest.TestCase):
    def test_unsupported_format(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            original = tmp_path / "original.txt"
            translated = tmp_path / "translated.txt"

            original.write_text("test")
            translated.write_text("test")

            with self.assertRaises(ValueError):
                validate_file_pair(original, translated)

    def test_format_mismatch(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)

            if DOCX_AVAILABLE:
                docx_file = tmp_path / "test.docx"
                _make_test_docx(docx_file)

            if OPENPYXL_AVAILABLE:
                xlsx_file = tmp_path / "test.xlsx"
                _make_test_xlsx(xlsx_file)

            if DOCX_AVAILABLE and OPENPYXL_AVAILABLE:
                with self.assertRaises(ValueError):
                    validate_file_pair(docx_file, xlsx_file)


class ValidateJobArtifactsTest(unittest.TestCase):
    def test_empty_lists(self) -> None:
        results = validate_job_artifacts(
            review_dir=Path("/tmp"),
            original_files=[],
            translated_files=[],
        )
        self.assertEqual(len(results), 0)


if __name__ == "__main__":
    unittest.main()
